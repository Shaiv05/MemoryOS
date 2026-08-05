"""
Document text extraction module.

Supports PDF (via PyMuPDF with pypdf fallback), DOCX, TXT/MD, images (OCR),
and web links.  Produces clean, structure-aware text with heading markers
(# / ## / ###) and page boundary markers (--- Page N ---).
"""

import logging
import re
import time
import unicodedata
from collections import Counter
from pathlib import Path

logger = logging.getLogger(__name__)


class ExtractionError(ValueError):
    pass


# ─── Unicode & Character Cleaning ───────────────────────────────────────────


def clean_raw_characters(text: str) -> str:
    """Removes null bytes, control characters, and normalizes Unicode."""
    if not text:
        return ""
    # Strip null characters and non-printable control characters
    # (preserve newline \n, tab \t, carriage return \r)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    # Unicode NFKC normalization (collapses compatibility characters)
    text = unicodedata.normalize("NFKC", text)
    return text


def remove_ocr_garbage(text: str) -> str:
    """Filters lines composed mostly of non-alphanumeric noise symbols."""
    clean_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            clean_lines.append("")
            continue
        # If line has more than 5 chars and over 65% are weird symbols, drop it
        if len(stripped) > 5:
            alpha_num = sum(1 for c in stripped if c.isalnum() or c.isspace())
            ratio = alpha_num / len(stripped)
            if ratio < 0.35:
                continue
        clean_lines.append(stripped)
    return "\n".join(clean_lines)


def dehyphenate_text(text: str) -> str:
    """Merges words broken across line breaks with a trailing hyphen."""
    if not text:
        return ""
    # Match word-hyphen-newline-word patterns
    text = re.sub(
        r"(\b[a-zA-Z]{2,})-\s*\n\s*([a-zA-Z]{2,}\b)",
        r"\1\2",
        text,
    )
    return text


def remove_page_numbers(text: str) -> str:
    """Remove standalone page number lines (e.g. '5', 'Page 5', '- 5 -')."""
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append(line)
            continue
        # Standalone number
        if re.match(r"^\d{1,4}$", stripped):
            continue
        # "Page 5" or "page 5"
        if re.match(r"^page\s+\d{1,4}$", stripped, re.IGNORECASE):
            continue
        # "- 5 -" or "— 5 —"
        if re.match(r"^[-–—]\s*\d{1,4}\s*[-–—]$", stripped):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def normalize_whitespace(text: str) -> str:
    """Collapse excessive whitespace while preserving paragraph structure."""
    if not text:
        return ""
    text = text.replace("\t", "    ")
    # Collapse runs of 3+ spaces to a single space
    text = re.sub(r"[ ]{3,}", " ", text)
    # Remove trailing whitespace per line
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
    # Collapse 4+ consecutive blank lines to double newline
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text


def merge_paragraph_lines(text: str) -> str:
    """
    Intelligently merges single line-breaks within paragraphs while preserving:
    - Double newlines (paragraph separators)
    - Bullet points / numbered lists
    - Section headings (# prefixed, ALL CAPS)
    - Page markers ('--- Page N ---')
    - Code-like indented blocks
    """
    if not text:
        return ""

    lines = text.splitlines()
    processed_blocks: list[str] = []
    current_block: list[str] = []

    list_pattern = re.compile(
        r"^\s*([*\-•▪◦‣]\s+|\d+[.)]\s+|[a-zA-Z][.)]\s+)"
    )
    heading_pattern = re.compile(r"^#{1,6}\s")
    page_marker_pattern = re.compile(r"^---\s*Page\s+\d+\s*---$", re.IGNORECASE)

    for line in lines:
        stripped = line.strip()

        # Empty line → flush current block, emit paragraph separator
        if not stripped:
            if current_block:
                processed_blocks.append(" ".join(current_block))
                current_block = []
            processed_blocks.append("")
            continue

        is_list = bool(list_pattern.match(stripped))
        is_heading = bool(heading_pattern.match(stripped)) or (
            len(stripped) < 100
            and stripped.isupper()
            and len(stripped) > 3
            and not stripped.startswith("---")
        )
        is_page_marker = bool(page_marker_pattern.match(stripped))

        if is_list or is_heading or is_page_marker:
            if current_block:
                processed_blocks.append(" ".join(current_block))
                current_block = []
            processed_blocks.append(stripped)
        else:
            if current_block:
                last_word = current_block[-1]
                # If previous block ends with sentence-ending punctuation,
                # start a new paragraph
                if last_word.endswith((".", ":", "!", "?", '."', ".'", '?"', "?'")):
                    processed_blocks.append(" ".join(current_block))
                    current_block = [stripped]
                else:
                    current_block.append(stripped)
            else:
                current_block.append(stripped)

    if current_block:
        processed_blocks.append(" ".join(current_block))

    # Collapse multiple consecutive empty lines to one
    final_output: list[str] = []
    prev_empty = False
    for block in processed_blocks:
        if not block:
            if not prev_empty:
                final_output.append("")
                prev_empty = True
        else:
            final_output.append(block)
            prev_empty = False

    return "\n".join(final_output).strip()


def fix_spaced_characters(text: str) -> str:
    """
    Fixes character kerning artifacts from PDFs where individual letters are separated by single spaces
    (e.g., 'd a t a   s t r u c t u r e' -> 'data structure', 'a d a t a' -> 'a data').
    """
    if not text:
        return ""

    # Match sequences where letters are separated by spaces, e.g. "a d a t a s t r u c t u r e"
    def _fix_kerning_block(match):
        raw = match.group(0)
        # Split on double spaces or newlines which represent real word/line breaks
        tokens = re.split(r"(\s{2,}|\n)", raw)
        fixed_tokens = []
        for tok in tokens:
            if tok.strip() and not re.match(r"^\s+$", tok):
                # Remove single spaces between single letters
                cleaned = re.sub(r"(?<=\b[a-zA-Z])\s+(?=[a-zA-Z]\b)", "", tok)
                fixed_tokens.append(cleaned)
            else:
                fixed_tokens.append(tok)
        return "".join(fixed_tokens)

    # Pattern matching 3 or more single-letter spaced tokens
    pattern = r"\b(?:[a-zA-Z]\s+){2,}[a-zA-Z]\b"
    text = re.sub(pattern, _fix_kerning_block, text)
    return text


def strip_explicit_header_footer_patterns(text: str) -> str:
    """
    Strips document header/footer artifacts such as:
    - 'Prepared By : Mayank Yadav'
    - Running page headers like 'Distributed System'
    - Unit closing lines like 'See you in the next unit !'
    """
    if not text:
        return ""

    header_patterns = [
        r"(?i)^\s*prepared\s+by\s*:.*$",
        r"(?i)^\s*see\s+you\s+in\s+the\s+next\s+unit\s*!.*$",
    ]

    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue

        # Skip explicit header patterns
        if any(re.match(pat, stripped) for pat in header_patterns):
            continue

        # Skip short floating header lines at the top of pages (e.g. standalone "Distributed System")
        if stripped.lower() in {"distributed system", "distributed systems"}:
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def normalize_text(text: str) -> str:
    """Full cleaning pipeline for extracted document text."""
    if not text:
        return ""
    text = clean_raw_characters(text)
    text = fix_spaced_characters(text)
    text = strip_explicit_header_footer_patterns(text)
    text = remove_ocr_garbage(text)
    text = dehyphenate_text(text)
    text = remove_page_numbers(text)
    text = merge_paragraph_lines(text)
    text = normalize_whitespace(text)
    return text.strip()


# ─── Header / Footer Detection ──────────────────────────────────────────────


def strip_repeated_headers_footers(page_texts: list[str]) -> list[str]:
    """Identifies and removes headers/footers repeated across PDF pages."""
    if len(page_texts) <= 2:
        return page_texts

    NUM_LINES = 2
    top_lines_per_page: list[list[str]] = []
    bottom_lines_per_page: list[list[str]] = []

    for pt in page_texts:
        lines = [ln.strip() for ln in pt.splitlines() if ln.strip()]
        top_lines_per_page.append(lines[:NUM_LINES] if len(lines) >= NUM_LINES else lines)
        bottom_lines_per_page.append(lines[-NUM_LINES:] if len(lines) >= NUM_LINES else lines)

    # Normalise numbers to 'N' so "Page 1" and "Page 2" match
    def _norm(s: str) -> str:
        return re.sub(r"\d+", "N", s)

    top_counter: Counter[str] = Counter()
    bottom_counter: Counter[str] = Counter()
    for lines in top_lines_per_page:
        for ln in lines:
            top_counter[_norm(ln)] += 1
    for lines in bottom_lines_per_page:
        for ln in lines:
            bottom_counter[_norm(ln)] += 1

    num_pages = len(page_texts)
    threshold = max(2, int(num_pages * 0.4))
    repeated_top = {pat for pat, cnt in top_counter.items() if cnt >= threshold}
    repeated_bottom = {pat for pat, cnt in bottom_counter.items() if cnt >= threshold}

    cleaned_pages = []
    for pt in page_texts:
        lines = pt.splitlines()
        filtered = []
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                filtered.append(line)
                continue
            normed = _norm(stripped)
            is_top = idx < NUM_LINES
            is_bottom = idx >= len(lines) - NUM_LINES
            if is_top and normed in repeated_top:
                continue
            if is_bottom and normed in repeated_bottom:
                continue
            filtered.append(line)
        cleaned_pages.append("\n".join(filtered))

    return cleaned_pages


# ─── PyMuPDF Extraction (Primary — Superior Quality) ────────────────────────


def _detect_body_font_size(doc) -> float:
    """Determine the dominant body-text font size by sampling pages."""
    try:
        import fitz  # noqa: F811
    except ImportError:
        return 12.0

    font_sizes: list[float] = []
    max_pages = min(len(doc), 15)

    for page_num in range(max_pages):
        page = doc[page_num]
        try:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        except Exception:
            continue
        for block in blocks:
            if block.get("type", 0) != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if len(text) > 10:
                        font_sizes.append(round(span.get("size", 12.0), 1))

    if not font_sizes:
        return 12.0
    return Counter(font_sizes).most_common(1)[0][0]


def _heading_level_from_span(span: dict, body_size: float) -> int:
    """
    Return heading level (1-4) if span looks like a heading, else 0.
    Based on font size ratio to body text and bold flag.
    """
    size = span.get("size", 12.0)
    flags = span.get("flags", 0)
    is_bold = bool(flags & (1 << 4))
    ratio = size / body_size if body_size > 0 else 1.0

    if ratio >= 1.8:
        return 1
    if ratio >= 1.4:
        return 2
    if ratio >= 1.15:
        return 3
    if is_bold and ratio >= 1.0:
        return 4
    return 0


def extract_pdf_text_pymupdf(file_path) -> list[str]:
    """Extract PDF text using PyMuPDF with structure-aware heading detection."""
    import fitz

    start = time.time()
    doc = fitz.open(str(file_path))
    body_size = _detect_body_font_size(doc)

    page_texts: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        try:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        except Exception:
            page_texts.append(page.get_text("text") or "")
            continue

        page_lines: list[str] = []

        for block in blocks:
            if block.get("type", 0) != 0:  # skip image blocks
                continue

            block_parts: list[str] = []
            block_heading_level = 0

            for line in block.get("lines", []):
                line_spans: list[str] = []
                line_max_level = 0

                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text:
                        continue
                    h_level = _heading_level_from_span(span, body_size)
                    if h_level > 0 and text.strip():
                        line_max_level = max(line_max_level, h_level)

                    # Ensure proper space between spans when font styles change
                    if line_spans:
                        prev = line_spans[-1]
                        if (
                            prev
                            and not prev[-1].isspace()
                            and not text[0].isspace()
                            and text[0] not in ".,;:!?')]}%-"
                        ):
                            line_spans.append(" ")
                    line_spans.append(text)

                line_text = "".join(line_spans).strip()
                if line_text:
                    block_heading_level = max(block_heading_level, line_max_level)
                    block_parts.append(line_text)

            if not block_parts:
                continue

            block_text = " ".join(block_parts)

            # Mark detected headings (but not overly-long blocks)
            if block_heading_level > 0 and len(block_text) < 200:
                prefix = "#" * min(block_heading_level, 4)
                if not block_text.startswith("#"):
                    block_text = f"{prefix} {block_text}"

            page_lines.append(block_text)

        page_texts.append("\n\n".join(page_lines))

    doc.close()
    logger.info(
        "PyMuPDF extraction: %d pages in %.2fs",
        len(page_texts),
        time.time() - start,
    )
    return page_texts


def extract_pdf_text_pypdf(file_path_or_obj) -> list[str]:
    """Fallback PDF extraction using pypdf."""
    from pypdf import PdfReader

    start = time.time()
    reader = PdfReader(file_path_or_obj)
    page_texts = [page.extract_text() or "" for page in reader.pages]
    logger.info(
        "pypdf extraction: %d pages in %.2fs",
        len(page_texts),
        time.time() - start,
    )
    return page_texts


def extract_pdf_text(file_path_or_obj) -> str:
    """
    Extract text from PDF.
    Tries PyMuPDF first (much better quality), falls back to pypdf.
    Returns clean text with page markers and heading markers.
    """
    page_raw_texts: list[str] = []
    extraction_method = "unknown"

    # --- Attempt PyMuPDF (primary) ---
    try:
        import fitz  # noqa: F401

        file_path = (
            file_path_or_obj
            if isinstance(file_path_or_obj, (str, Path))
            else getattr(file_path_or_obj, "path", file_path_or_obj)
        )
        page_raw_texts = extract_pdf_text_pymupdf(file_path)
        extraction_method = "pymupdf"
    except ImportError:
        logger.info("PyMuPDF not installed — falling back to pypdf.")
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed (%s) — trying pypdf.", exc)

    # --- Fallback to pypdf ---
    if not page_raw_texts or all(not p.strip() for p in page_raw_texts):
        try:
            page_raw_texts = extract_pdf_text_pypdf(file_path_or_obj)
            extraction_method = "pypdf"
        except Exception as exc:
            raise ExtractionError(f"Failed to read PDF file: {exc}") from exc

    if not page_raw_texts:
        raise ExtractionError("No pages found in PDF.")

    # --- Clean repeated headers / footers ---
    cleaned_pages = strip_repeated_headers_footers(page_raw_texts)

    # --- Assemble with page markers ---
    text_parts: list[str] = []
    for page_idx, page_text in enumerate(cleaned_pages):
        cleaned = page_text.strip()
        if cleaned:
            text_parts.append(f"--- Page {page_idx + 1} ---\n{cleaned}")

    text = normalize_text("\n\n".join(text_parts))

    if not text:
        raise ExtractionError("No selectable text found in this PDF.")

    logger.info(
        "PDF extraction complete via %s: %d chars, %d pages",
        extraction_method,
        len(text),
        len(cleaned_pages),
    )
    return text


# ─── TXT / Markdown ─────────────────────────────────────────────────────────


def extract_txt_text(file_obj) -> str:
    try:
        if hasattr(file_obj, "open"):
            file_obj.open("rb")
            raw = file_obj.read()
        elif hasattr(file_obj, "read"):
            raw = file_obj.read()
        else:
            with open(file_obj, "rb") as f:
                raw = f.read()
    except Exception:
        if hasattr(file_obj, "read"):
            raw = file_obj.read()
        else:
            with open(file_obj, "rb") as f:
                raw = f.read()
    finally:
        try:
            if hasattr(file_obj, "close"):
                file_obj.close()
        except Exception:
            pass

    if isinstance(raw, bytes):
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="replace")
    else:
        text = str(raw)

    text = normalize_text(text)
    if not text:
        raise ExtractionError("Text file is empty.")
    return text


# ─── DOCX ────────────────────────────────────────────────────────────────────


def extract_docx_text(file_path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError(
            "python-docx is not installed. Run pip install python-docx."
        ) from exc

    try:
        doc = docx.Document(file_path)
        text_parts: list[str] = []

        for paragraph in doc.paragraphs:
            p_text = paragraph.text.strip()
            if not p_text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                level = 1
                try:
                    level = int(paragraph.style.name.replace("Heading", "").strip())
                except (ValueError, TypeError):
                    level = 2
                prefix = "#" * min(level, 4)
                text_parts.append(f"{prefix} {p_text}")
            else:
                text_parts.append(p_text)

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                text_parts.append("\n" + "\n".join(table_rows) + "\n")

        text = normalize_text("\n\n".join(text_parts))
    except Exception as exc:
        raise ExtractionError(f"Failed to extract DOCX text: {exc}") from exc

    if not text:
        raise ExtractionError("DOCX file contains no readable text.")
    return text


# ─── Image (OCR) ─────────────────────────────────────────────────────────────


def extract_image_text(file_path) -> str:
    """Attempts OCR extraction using pytesseract/PIL, falls back gracefully."""
    try:
        from PIL import Image
    except ImportError:
        raise ExtractionError(
            "PIL library is missing. Install Pillow to process images."
        )

    try:
        img = Image.open(file_path)
        try:
            import pytesseract

            ocr_text = pytesseract.image_to_string(img)
            text = normalize_text(ocr_text)
            if text:
                return text
        except Exception:
            pass

        return normalize_text(
            f"Image File: {Path(file_path).name}\n"
            f"Format: {img.format}, Size: {img.size[0]}x{img.size[1]}px, Mode: {img.mode}\n"
            "Visual image document ingested into MemoryOS."
        )
    except Exception as exc:
        raise ExtractionError(f"Failed to process image file: {exc}") from exc


# ─── Link / URL ──────────────────────────────────────────────────────────────


def extract_link_text(url: str) -> str:
    """Fetches web page content and converts HTML to clean text."""
    if not url:
        raise ExtractionError("Link URL is missing.")

    try:
        import requests

        response = requests.get(
            url, timeout=10, headers={"User-Agent": "MemoryOS/1.0"}
        )
        response.raise_for_status()

        html = response.text
        clean_text = re.sub(
            r"<script.*?>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE
        )
        clean_text = re.sub(
            r"<style.*?>.*?</style>", "", clean_text, flags=re.DOTALL | re.IGNORECASE
        )
        clean_text = re.sub(r"<[^>]+>", " ", clean_text)
        clean_text = re.sub(r"\s+", " ", clean_text)

        text = normalize_text(clean_text)
        if not text:
            raise ExtractionError(f"No text content found at URL: {url}")
        return text
    except Exception as exc:
        raise ExtractionError(
            f"Failed to fetch content from URL ({url}): {exc}"
        ) from exc


# ─── Dispatcher ──────────────────────────────────────────────────────────────


def extract_document_text(document) -> str:
    """Extract text from a Document model instance based on its file_type."""
    if document.file_type == "note":
        text = normalize_text(document.raw_text or "")
        if not text:
            raise ExtractionError("Note content is empty.")
        return text

    if document.file_type == "pdf":
        return extract_pdf_text(document.file.path)

    if document.file_type in {"txt", "md"}:
        return extract_txt_text(document.file)

    if document.file_type == "docx":
        return extract_docx_text(document.file.path)

    if document.file_type == "image":
        return extract_image_text(document.file.path)

    if document.file_type == "link":
        if document.raw_text and document.raw_text.strip():
            return normalize_text(document.raw_text)
        return extract_link_text(document.source_url)

    raise ExtractionError(f"Unsupported document type: {document.file_type}")
